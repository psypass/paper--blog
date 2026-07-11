from __future__ import annotations

from pathlib import Path

from docx import Document

from paper_blog_agent.ingestion.common import extract_abstract, extract_authors, file_sha256, first_nonempty_line
from paper_blog_agent.models import PaperSource


class DocxLoader:
    def load(self, path: str) -> PaperSource:
        file_path = Path(path)
        document = Document(str(file_path))
        paragraphs = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))
        text = "\n\n".join(paragraphs)
        return PaperSource(
            source_type="docx",
            source_id=file_sha256(file_path),
            title=first_nonempty_line(text, file_path.stem),
            authors=extract_authors(text),
            abstract=extract_abstract(text),
            raw_text=text,
            metadata={"filename": file_path.name},
            original_path=str(file_path),
        )
